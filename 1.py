#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Date    : 2019-09-22 12:37:03
# @Author  : Zhu, Yaohui (yaohui.zhu@beigene.com)
# @Link    : https://github.com/zyhwhu2009/
# @Version : $Id$

import os,copy
# from docx import *

from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import docx

# file = docx.Document('test.docx')

# print(len(file.paragraphs))

# print(len(file.tables))

# for it in file.tables:
#     # print(it)

#     for row in it.rows:

#         for cell in row.cells:
#             # print( cell.text )
#             pass


# table_select    = file.tables[1]

# rows=table_select.rows


# print("Table rows: "+str(len(rows)))

# cols=rows[0].cells


# print('Table columns :'+ str(len(cols)))
# cell=cols[0]
# print(cell.text)

# style  = copy.deepcopy(table_select.style)
# print(style)

# rows_len=len(rows)

# cols_len=len(cols)


# ///////////////////////////////////////////////////////////
# 将表格复制到新的文档

# file_new = Document()

# # 添加表格
# table = file_new.add_table(rows=rows_len, cols=cols_len,style=style)
# # 添加表格内容，也可用： table.rows[0].cells[0].text = "第一行第一列"
# #               或者用： table.cell(0,0).text = "cell_00"
# for i in range(rows_len):
#     for j in range(cols_len):
#         cell = table.cell(i, j)
#         # cell.text = "第"+str(i+1) +"行第"+str(j+1) +"列"
#         cell.text=table_select.cell(i,j).text
# file_new.save('1.docx')
# //////////////////////////////////////////////////////////////////


# ///////////////////////////////////////////////////////////////////
# 搜索特定的字符




# def search_word(filename,word):
#     #打开文档
#     # document = Document(filename)
#     document = filename
#     # document = Document(r'C:\Users\Cheng\Desktop\kword\words\wind.docx')
#     #读取每段资料
#     l = [ paragraph for paragraph in document.paragraphs];

#     #输出并观察结果，也可以通过其他手段处理文本即可
#     for para in l:

#         i=para.text.strip()
#         # print i

#         if i.find(word)!=-1:
#             print(i)
#             para.runs[-1].add_break(enum.text.WD_BREAK.PAGE)


# search_word(file,'表2')


# file.save('example4.docx')


from  copy import *

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    # print(parent_elm)
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield [Paragraph(child, parent),1]
        elif isinstance(child, CT_Tbl):
            yield [Table(child, parent),0]
# import docx
# test=iter_block_items(file)
doc = docx.Document('1.docx')

# print(iter_block_items(doc))
obj      =   [block for block in iter_block_items(doc)]

print(obj[0][1])
n=1


def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)

    p.addnext(new_tbl)
    p.addnext(new_tbl)


for block in iter_block_items(doc):
    # print(block[0])
    try:
        if block[0].text.lower().startswith('test'):
            # print(str(n),block[0].text)
            table_before_n = 0
            for i in range(n,-1,-1):
                # print(obj[i][1])
                if obj[i][1]  == 0:
                    print('dasdasd',str(i))
                    table_before_n = i
                    break
            table_to_copy = obj[table_before_n][0]

            copy_table_after(table_to_copy, block[0])
            block[0].text = '1' + block[0].text
            new_p=block[0].insert_paragraph_before ('end_of_table')
            new_p.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

        # print(str(n),block[0].text)
    except Exception as e:
        # raise
        print(e)
        print(str(n),len(block[0].rows))
    n=n+1


doc.save('2.docx')

